"""Generate Workflowiga blocks & configuration reference as a Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Workflowiga-Blocks-Configuration-Reference.docx"


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F97316")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def add_block(
    doc: Document,
    name: str,
    category: str,
    policy_types: str,
    description: str,
    config_rows: list[list[str]],
    notes: str | None = None,
) -> None:
    doc.add_heading(name, level=2)
    meta = doc.add_paragraph()
    meta.add_run("Category: ").bold = True
    meta.add_run(f"{category}    ")
    meta.add_run("Policy types: ").bold = True
    meta.add_run(policy_types)
    doc.add_paragraph(description)
    doc.add_heading("Configuration", level=3)
    add_table(doc, ["Field / Section", "Type", "Options / Details"], config_rows)
    if notes:
        p = doc.add_paragraph()
        p.add_run("Notes: ").bold = True
        p.add_run(notes)
    doc.add_paragraph()


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Title
    title = doc.add_heading("Workflowiga — Blocks & Configuration Reference", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Reference guide for all canvas blocks and their configuration panels "
        "in Workflow Policies and Approval Policies."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Generated for Workflowiga IAM workflow builder.")
    doc.add_page_break()

    # Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Workflowiga supports two policy types, each with its own editor context:"
    )
    add_table(
        doc,
        ["Policy type", "Trigger event", "Typical use"],
        [
            [
                "Workflow Policy",
                "Joiner (Mover / Leaver coming soon)",
                "Lifecycle automation — provision access, run linked approval policies, branch by attributes.",
            ],
            [
                "Approval Policy",
                "Approval Policy",
                "Multi-step sign-off — approvers, SoD checks, notifications, conditional routing.",
            ],
        ],
    )
    doc.add_paragraph(
        "Blocks are dragged from the left palette onto the canvas. Clicking a block "
        "opens the configuration panel on the right. Block names are edited in the "
        "panel header (click the title). Changes apply live — there is no separate Apply button."
    )
    doc.add_paragraph(
        "Palette categories: Events, Filters (workflow only), Tasks, Modules (workflow only), Rules."
    )

    # Workflow policies
    doc.add_page_break()
    doc.add_heading("2. Workflow Policy Blocks", level=1)
    doc.add_paragraph(
        "Available when the policy event is Joiner, Mover, or Leaver. "
        "Canvas order: Start → Event → User Filter (optional) → Tasks → End."
    )

    add_block(
        doc,
        "Joiner (Event)",
        "Events",
        "Workflow only",
        "Lifecycle trigger fired when a new identity is created. Required first block on the canvas.",
        [
            ["Description", "Textarea", "Optional text shown on the event card."],
        ],
    )

    add_block(
        doc,
        "Mover (Event)",
        "Events",
        "Workflow only (coming soon)",
        "Trigger when an identity's attributes change. Currently marked coming soon in the palette.",
        [["Description", "Textarea", "Optional text shown on the event card."]],
    )

    add_block(
        doc,
        "Leaver (Event)",
        "Events",
        "Workflow only (coming soon)",
        "Trigger when an identity is deactivated. Currently marked coming soon in the palette.",
        [["Description", "Textarea", "Optional text shown on the event card."]],
    )

    add_block(
        doc,
        "User Filter",
        "Filters",
        "Workflow only",
        "Scopes the workflow to users matching attribute-based conditions. Placed below the event.",
        [
            ["Define User Conditions", "Drawer → Condition builder", "AND/OR logic; up to 10 conditions per rule."],
            ["Condition row", "Attribute + Operator + Value", "Identity attributes (Department, Location, Grade, etc.)."],
            ["Operators", "Dropdown", "equals, not equals, contains, starts with, in, not in, greater/less than, gte, lte."],
        ],
        notes="Opens a side drawer with the full ConditionBuilder. Summary chips shown on the config card.",
    )

    add_block(
        doc,
        "Assign Entities",
        "Tasks",
        "Workflow only (also embeddable in branch columns)",
        "Provisions applications, entitlements, technical roles, and business roles to matched users.",
        [
            ["Block name", "Header (editable)", "Displayed on canvas card."],
            ["Add Apps & Entitlements", "Drawer picker", "Select applications, then refine with entitlements."],
            ["Add Technical Roles", "Drawer picker", "Reusable entitlement bundles across apps."],
            ["Add Business Roles", "Drawer picker", "Organisation role bundles."],
            ["Apply Assignment Criteria", "Drawer → Condition builder", "Optional AND/OR rules to further narrow who receives the assignment within the filter."],
        ],
    )

    add_block(
        doc,
        "Notification",
        "Tasks",
        "Workflow & Approval (top-level or embedded in branches)",
        "Sends Slack and/or email when a preceding step starts, completes, or fails.",
        [
            ["When to send", "Segmented control", "On start | On complete | On failure (hidden when embedded in a branch)."],
            ["Audience", "Multi-select", "Requester, Manager, Owner, Specific users."],
            ["Specific users", "Drawer picker", "Required when Specific users is selected."],
            ["Channels", "Modal (Configure channels)", "Enable Slack/Email; edit per-channel message with {{variable}} tokens and preview."],
        ],
        notes="Embedded branch notifications inherit trigger from branch context (no When to send section).",
    )

    add_block(
        doc,
        "Approval Policy (Module)",
        "Modules",
        "Workflow only",
        "Embeds an active approval policy from the library at this step in the workflow.",
        [
            ["Block name", "Header (editable)", "Label on canvas."],
            ["Select policy", "Drawer picker (single)", "Only active approval policies are selectable."],
            ["Open this policy", "Link button", "Opens the linked policy in the editor."],
        ],
    )

    # Rules shared in workflow
    doc.add_heading("Rules (also in Workflow policies)", level=2)

    add_block(
        doc,
        "Multisplit Branch",
        "Rules",
        "Workflow & Approval",
        "Splits flow into parallel branches by attribute values. Each branch can contain its own nested blocks.",
        [
            ["Split by attributes", "Multi-select", "One or more attributes (Department, Risk score, etc.)."],
            ["Branch definitions", "Per-branch dropdowns", "When 1+ attributes selected: set value per branch; add/remove branches."],
            ["Branches count", "Number", "2–8 (auto when attributes drive branch count)."],
            ["Completion rule", "Dropdown", "Approval only: All | Any | Majority | Threshold."],
            ["Threshold", "Number", "When completion rule is Threshold — N of M branches."],
            ["Global fallback type", "Dropdown", "Approval only: Skip | Block | Add fallback email."],
            ["Global fallback approvers", "Drawer picker", "When fallback is Add fallback email."],
            ["Global SLA", "Toggle + duration + unit", "Approval only: hours/days; On SLA timeout action."],
        ],
        notes="In workflow context, all branches run in parallel. Branch columns support nested Approval Level, Notification, Assign Entities, Conditional Branch, Skip, Exit.",
    )

    add_block(
        doc,
        "Conditional Branch",
        "Rules",
        "Workflow & Approval",
        "Routes the request down the first matching IF / ELSE IF branch; last branch is always ELSE (catch-all).",
        [
            ["Branch count", "Number", "2–8 branches."],
            ["Routing ladder", "Per-branch UI", "IF / ELSE IF / ELSE rows with editable route names."],
            ["Conditions", "Drawer → Condition builder", "Per IF/ELSE IF branch via sliders icon; attributes from routing set."],
            ["Branch Defaults", "Section", "Approval only: global fallback type, fallback users, global SLA settings."],
        ],
        notes="ELSE branch always shows 'All other requests' — no condition editor. First match wins.",
    )

    add_block(
        doc,
        "Skip",
        "Rules",
        "Workflow & Approval",
        "Bypasses steps on a branch route; flow continues after the branch merge.",
        [["—", "—", "No configuration. Informational panel only."]],
    )

    add_block(
        doc,
        "Exit",
        "Rules",
        "Workflow & Approval",
        "Terminates a branch or top-level flow. No connector is drawn below an Exit block.",
        [["—", "—", "No configuration. Informational panel only."]],
    )

    # Approval policies
    doc.add_page_break()
    doc.add_heading("3. Approval Policy Blocks", level=1)
    doc.add_paragraph(
        "Available when the policy event is Approval Policy. "
        "Palette groups: Tasks and Rules."
    )

    add_block(
        doc,
        "Approval Policy (Event)",
        "Events",
        "Approval only",
        "Root event for an approval policy. Describes what the policy governs.",
        [["Description", "Textarea", "Optional text shown on the policy event card."]],
    )

    add_block(
        doc,
        "Approval Level",
        "Tasks",
        "Approval only (also embeddable in branch columns)",
        "Single approver step with fallback, SLA, and optional assignment alerts.",
        [
            ["Block name", "Header (editable)", "Shown as card title on canvas."],
            ["Approver type", "Dropdown", "Manager | Owner | Governance Group | User | Custom attribute."],
            ["Select Governance Group / User", "Drawer picker", "When approver type is entity-based; multi-select."],
            ["Resolve approver from", "Drawer", "When Custom attribute: pick resolver (Dept Head, App Owner, etc.) + optional Apply when conditions."],
            ["New request alert", "Toggle + Channels modal", "Notify approver on assignment via Slack/Email with custom message."],
            ["Fallback type", "Dropdown", "Skip | Block | Add fallback email. Override toggle when inside Multisplit."],
            ["Fallback approvers", "Drawer picker", "Users when fallback is Add fallback email."],
            ["Enable SLA", "Toggle", "Override toggle when inside Multisplit."],
            ["SLA duration + unit", "Number + dropdown", "hours | days."],
            ["On SLA timeout", "Dropdown", "Auto-escalate to manager | Notify Only | Auto Approve | Auto Reject."],
        ],
        notes="Approver type pill shown on canvas card. Configured chip when approver and fallback are set.",
    )

    add_block(
        doc,
        "SoD Check",
        "Tasks",
        "Approval only",
        "Segregation of Duties pre-check evaluated before downstream approval steps.",
        [
            ["Block name", "Header (editable)", "Label on canvas."],
            ["If SoD violation detected → Action", "Dropdown", "Continue to next step | Send notification only | End flow."],
            ["Also notify / Channels", "Modal", "When Continue or Notify: configure Slack/Email templates."],
            ["Notify", "Multi-select", "Requester, Manager, Owner, Specific users."],
            ["Specific users", "Drawer picker", "When Specific users selected."],
            ["If no SoD violation", "Info", "Request continues automatically — no config."],
        ],
    )

    doc.add_paragraph(
        "Notification, Multisplit Branch, Conditional Branch, Skip, and Exit blocks "
        "in approval policies use the same configuration as described in Section 2."
    )

    # Branch embedded
    doc.add_page_break()
    doc.add_heading("4. Branch-Embedded Blocks", level=1)
    doc.add_paragraph(
        "Inside Multisplit Branch or Conditional Branch columns, these block types can be nested:"
    )
    add_table(
        doc,
        ["Block", "Embedded behaviour"],
        [
            ["Approval Level", "Same config; can override parent global Fallback/SLA."],
            ["Notification", "No 'When to send' — trigger inherited from branch context."],
            ["Assign Entities", "Workflow branches only — same assignment drawers."],
            ["Conditional Branch", "Nested routing (embedded conditional) inside a column."],
            ["Skip / Exit", "Terminal markers within the branch column."],
        ],
    )

    doc.add_heading("Conditional Branch — per-branch condition attributes", level=2)
    doc.add_paragraph(
        "Each IF/ELSE IF branch opens a drawer with ConditionBuilder. Available attributes include:"
    )
    attrs = (
        "Routing signals: Amount, Risk score, Entitlement count. "
        "Access context: Policy type, Business role, User access risk score, Has baseline access, "
        "Access decision, Owner decision, SoD violation, Requested application, entitlement fields. "
        "Identity: Department, Location, Grade, Employee Type, Manager, etc."
    )
    doc.add_paragraph(attrs)

    # Common UI patterns
    doc.add_heading("5. Common Configuration Patterns", level=1)
    add_table(
        doc,
        ["Pattern", "Used in"],
        [
            ["SectionCard → Drawer", "Entity pickers (users, groups, policies), condition builders, custom approver."],
            ["ChannelsField → Modal", "Notification & SoD violation alerts — Slack/Email toggle + message editor + variables."],
            ["ConfigInset / ConfigRow", "Grouped settings with grey background (e.g. New request alert)."],
            ["Live apply", "All panels — no footer Apply/Close; canvas updates immediately."],
            ["EditableName in header", "All task blocks — click title in config panel header to rename."],
        ],
    )

    doc.add_heading("6. Notification audiences & channels", level=1)
    add_table(
        doc,
        ["Audience", "Description"],
        [
            ["Requester", "Person who submitted the access request."],
            ["Manager", "Requester's manager."],
            ["Owner", "Application or entitlement owner."],
            ["Specific users", "Manually picked users from directory."],
        ],
    )
    doc.add_paragraph("Channels: Slack and Email. Each channel supports a custom message template with variables such as {{requester.name}}, {{application.name}}, {{entitlement.name}}, {{policy.name}}, {{status}}.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
